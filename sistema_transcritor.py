import os
import zipfile
from google.cloud import vision
import io
from docx import Document

# 1. AUTENTICAÇÃO CORRIGIDA
# O cliente será instanciado sem argumentos.
# Ele vai procurar automaticamente a variável de ambiente
# GOOGLE_APPLICATION_CREDENTIALS que você definiu no PowerShell.
try:
    client = vision.ImageAnnotatorClient()
    print("Cliente Vision API autenticado com sucesso.")
except Exception as e:
    print(f"Erro ao autenticar: {e}")
    print("Verifique se você definiu a variável GOOGLE_APPLICATION_CREDENTIALS no seu terminal.")
    exit()


# Função para descompactar o arquivo ZIP
def unzip_images(zip_file_path, extract_folder):
    # Garante que a pasta de extração exista
    if not os.path.exists(extract_folder):
        os.makedirs(extract_folder)
        
    with zipfile.ZipFile(zip_file_path, 'r') as zip_ref:
        zip_ref.extractall(extract_folder)
        
    extracted_files = os.listdir(extract_folder)
    extracted_files.sort()  # Ordena as imagens
    return extracted_files

# Função para realizar OCR em uma imagem
def perform_ocr(image_path):
    with io.open(image_path, 'rb') as image_file:
        content = image_file.read()
    
    # Usando vision.Image em vez de types.Image (mais moderno)
    image = vision.Image(content=content)
    
    # Usando document_text_detection para texto denso (melhor para livros)
    response = client.document_text_detection(image=image)

    # Verificando se há erros no OCR
    if response.error.message:
        raise Exception(f'Error during OCR: {response.error.message}')
    
    # 4. CORREÇÃO DO INDEXERROR
    # Verifica se algum texto foi encontrado antes de tentar acessá-lo
    if response.full_text_annotation:
        return response.full_text_annotation.text
    else:
        return "" # Retorna string vazia se nada for encontrado

# Função para criar e salvar o documento DOCX
def save_text_to_docx(extracted_text, output_path):
    doc = Document()
    doc.add_heading('Transcrição do Livro: Cultura Surda', 0)
    
    for page, text in extracted_text.items():
        doc.add_heading(f'Página: {page}', level=2) # Use um heading menor para páginas
        doc.add_paragraph(text)
        doc.add_page_break() # Melhor que um separador
        
    doc.save(output_path)
    print(f"Documento salvo em: {output_path}")

# --- CONFIGURAÇÃO DE CAMINHOS ---
# 2. e 3. CORREÇÃO DE CAMINHOS
# Estamos usando caminhos relativos. O script espera que
# o .zip esteja na MESMA pasta que o script .py

# ATUALIZE ESTE NOME para o nome real do seu arquivo zip
zip_file_path = "meu_livro.zip" 
extracted_folder = "imagens_extraidas"
output_file_path = "Transcricao_Cultura_Surda.docx"

# --- EXECUÇÃO PRINCIPAL ---
try:
    # Descompacta o ZIP e obtém as imagens
    print(f"Descompactando '{zip_file_path}'...")
    extracted_files = unzip_images(zip_file_path, extracted_folder)
    print(f"{len(extracted_files)} imagens encontradas.")   

    # Processando as imagens e realizando OCR
    extracted_text = {}
    for image_file in extracted_files:
        # Garante que estamos processando apenas arquivos de imagem
        if not image_file.lower().endswith(('.png', '.jpg', '.jpeg', '.tiff', '.bmp')):
            continue

        image_path = os.path.join(extracted_folder, image_file)
        try:
            print(f'Processando a imagem: {image_file}')
            text = perform_ocr(image_path)
            extracted_text[image_file] = text
        except Exception as e:
            print(f'Erro ao processar a imagem {image_file}: {str(e)}')

    # Salvar o resultado em um documento DOCX
    save_text_to_docx(extracted_text, output_file_path)

except FileNotFoundError:
    print(f"ERRO: O arquivo '{zip_file_path}' não foi encontrado.")
    print("Verifique se o nome do arquivo está correto e se ele está na mesma pasta do script.")
except Exception as e:
    print(f"Um erro inesperado ocorreu: {e}")